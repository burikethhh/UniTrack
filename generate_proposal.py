"""
UniTrack Project Proposal Document Generator
Generates a complete Word document for Sultan Kudarat State University
With UI Mockups included
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Colors
WHITE = (255, 255, 255)
BLACK = (0, 0, 0)
DARK_BLUE = (0, 51, 102)
LIGHT_BLUE = (230, 242, 255)
GREEN = (0, 153, 76)
LIGHT_GREEN = (220, 255, 220)
GRAY = (128, 128, 128)
LIGHT_GRAY = (240, 240, 240)
RED = (220, 53, 69)
ORANGE = (255, 165, 0)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    """Set paragraph spacing"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def create_styled_heading(doc, text, level=1):
    """Create a styled heading with custom colors"""
    heading = doc.add_heading(text, level=level)
    
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 76, 153)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(51, 51, 51)
    
    set_paragraph_spacing(heading, before=12, after=6)
    return heading

def add_bullet_point(doc, text, bold_prefix=None):
    """Add a bullet point with optional bold prefix"""
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        run2 = para.add_run(text)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(para, before=2, after=4, line_spacing=1.15)
    return para

def add_numbered_item(doc, number, text):
    """Add a numbered list item"""
    para = doc.add_paragraph()
    run = para.add_run(f"{number}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run2 = para.add_run(text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(para, before=2, after=4, line_spacing=1.15)
    para.paragraph_format.left_indent = Inches(0.25)
    return para

def create_styled_table(doc, headers, data, header_color='003366'):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_data)
            if i % 2 == 0:
                set_cell_shading(cell, 'E8F4FC')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
    
    return table

def add_mockup_image(doc, image_path, title, width=2.5):
    """Add a mockup image with title"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)
        set_paragraph_spacing(caption, before=4, after=12)

def create_proposal():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # ==================== TITLE PAGE ====================
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SULTAN KUDARAT STATE UNIVERSITY")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = location.add_run("ACCESS, EJC Montilla, Tacurong City, Sultan Kudarat")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    for _ in range(4):
        doc.add_paragraph()
    
    project_title = doc.add_paragraph()
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_title.add_run("PROJECT PROPOSAL")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    app_name = doc.add_paragraph()
    app_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = app_name.add_run("UniTrack")
    run.bold = True
    run.font.size = Pt(36)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Real-Time Faculty & Staff Locator")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    doc.add_paragraph()
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("A Mobile-Based Geographic Information System\nfor Sultan Kudarat State University")
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    for _ in range(5):
        doc.add_paragraph()
    
    # Proposer info
    proposed_label = doc.add_paragraph()
    proposed_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proposed_label.add_run("Proposed by:")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    proposer = doc.add_paragraph()
    proposer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proposer.add_run("CHRISTIAN KETH AGUACITO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    for _ in range(2):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("February 2026")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    set_paragraph_spacing(toc_heading, before=0, after=24)
    
    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. Project Overview", "4"),
        ("3. Problem Statement", "5"),
        ("4. Objectives", "6"),
        ("5. Scope and Limitations", "7"),
        ("6. Proposed Features", "8"),
        ("7. Application User Interface Design", "10"),
        ("8. Technical Stack (Free & Open-Source)", "14"),
        ("9. System Architecture", "15"),
        ("10. Privacy & Security Measures", "16"),
        ("11. Development Roadmap", "17"),
        ("12. Project Timeline", "18"),
        ("13. Cost Analysis: Zero-Budget Implementation", "19"),
        ("14. Risk Assessment and Mitigation", "20"),
        ("15. Expected Outcomes", "21"),
        ("16. Evaluation Criteria", "22"),
        ("17. Future Expansion: IoT & Kiosk Integration", "23"),
        ("18. Conclusion", "25"),
        ("19. References", "26"),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run1 = p.add_run(item)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.color.rgb = RGBColor(0, 0, 0)
        run2 = p.add_run("." * (55 - len(item)))
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(180, 180, 180)
        run3 = p.add_run(page)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(12)
        run3.font.color.rgb = RGBColor(0, 51, 102)
        run3.bold = True
        set_paragraph_spacing(p, before=4, after=4, line_spacing=1.5)
    
    doc.add_page_break()
    
    # ==================== 1. EXECUTIVE SUMMARY ====================
    create_styled_heading(doc, "1. Executive Summary", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 128, 0)
    run2 = para.add_run(" is an innovative mobile-based Geographic Information System (GIS) designed specifically for Sultan Kudarat State University (SKSU) to enhance campus communication and efficiency. The application addresses a persistent challenge faced by students: locating faculty members and staff in real-time across the university's multiple buildings and facilities.")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    para2 = doc.add_paragraph()
    run = para2.add_run("The system employs a 'Google Maps' style interface powered by ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = para2.add_run("OpenStreetMap")
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 128, 0)
    run3 = para2.add_run(" (completely free and open-source) that allows students to view the real-time locations of teachers and staff within campus boundaries, provided the staff member has explicitly granted permission through a privacy-first consent mechanism.")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    set_paragraph_spacing(para2, before=0, after=12, line_spacing=1.5)
    para2.paragraph_format.first_line_indent = Inches(0.5)
    
    para3 = doc.add_paragraph()
    run = para3.add_run("Key features include a ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = para3.add_run("Staff Module (the 'Beacon')")
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run3 = para3.add_run(" with privacy controls and status presets, a ")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    run4 = para3.add_run("Student Module (the 'Seeker')")
    run4.bold = True
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)
    run5 = para3.add_run(" with live map visualization and navigation, and an ")
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(12)
    run6 = para3.add_run("Administrative Module")
    run6.bold = True
    run6.font.name = 'Times New Roman'
    run6.font.size = Pt(12)
    run7 = para3.add_run(" for department management.")
    run7.font.name = 'Times New Roman'
    run7.font.size = Pt(12)
    set_paragraph_spacing(para3, before=0, after=12, line_spacing=1.5)
    para3.paragraph_format.first_line_indent = Inches(0.5)
    
    highlight = doc.add_paragraph()
    highlight.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlight.add_run("TOTAL PROJECT COST: PHP 0.00 (Completely Free)")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)
    set_paragraph_spacing(highlight, before=12, after=12)
    
    para4 = doc.add_paragraph()
    run = para4.add_run("The system is built entirely on free and open-source technologies including Flutter, OpenStreetMap, and Firebase's free tier—resulting in zero development and operational costs. UniTrack represents SKSU's commitment to digital transformation and creating a more connected, efficient campus environment while maintaining the highest standards of privacy and data protection.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para4, before=0, after=12, line_spacing=1.5)
    para4.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== 2. PROJECT OVERVIEW ====================
    create_styled_heading(doc, "2. Project Overview", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack is a mobile-based geographic information system (GIS) designed to bridge the communication gap between students and university personnel at Sultan Kudarat State University. The application provides a live 'Google Maps' style interface where students can view the real-time location of teachers and staff within campus bounds, provided the staff member has granted permission.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    create_styled_heading(doc, "2.1 Project Title", level=2)
    p = doc.add_paragraph()
    run = p.add_run("UniTrack – Real-Time Faculty & Staff Locator for Sultan Kudarat State University")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    
    create_styled_heading(doc, "2.2 Project Type", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Mobile Application Development (Android & iOS) with Real-Time GIS Integration")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    
    create_styled_heading(doc, "2.3 Target Users", level=2)
    users = ["Students of Sultan Kudarat State University (Primary Users)", "Faculty Members and Teaching Staff", "Administrative Staff and Personnel", "University Administrators and Department Heads"]
    for user in users:
        add_bullet_point(doc, user)
    
    create_styled_heading(doc, "2.4 Target Deployment", level=2)
    p = doc.add_paragraph()
    run = p.add_run("The system will be deployed across all SKSU campuses, initially piloting at the Main Campus (ACCESS, EJC Montilla, Tacurong City) before expanding to satellite campuses.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    
    doc.add_page_break()
    
    # ==================== 3. PROBLEM STATEMENT ====================
    create_styled_heading(doc, "3. Problem Statement", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("Students at Sultan Kudarat State University often waste significant time walking to faculty offices only to find that the teacher is in a meeting, conducting a class in a different building, or temporarily off-campus. The current methods of locating faculty members present several challenges:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    problems = [
        ("Inefficient Communication Channels: ", "Email responses are often delayed, and physical bulletin boards with faculty schedules are frequently outdated or not accessible in real-time."),
        ("Time Wastage: ", "Students spend considerable time moving between buildings searching for available faculty, reducing productive academic hours."),
        ("Missed Consultation Opportunities: ", "Students miss valuable consultation time due to inability to locate faculty during their available hours."),
        ("Physical Traffic Congestion: ", "Faculty rooms become overcrowded during peak hours as students congregate hoping to catch their professors."),
        ("Lack of Real-Time Information: ", "No existing system provides live updates on faculty availability and location within the campus.")
    ]
    
    for title, desc in problems:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    doc.add_page_break()
    
    # ==================== 4. OBJECTIVES ====================
    create_styled_heading(doc, "4. Objectives", level=1)
    
    create_styled_heading(doc, "4.1 General Objective", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("To develop and implement a mobile-based real-time faculty and staff location system for Sultan Kudarat State University that enhances campus communication efficiency while maintaining strict privacy controls—all at zero cost using free and open-source technologies.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    create_styled_heading(doc, "4.2 Specific Objectives", level=2)
    
    objectives = [
        "To implement a live map interface using OpenStreetMap (free, open-source) for real-time visualization of faculty locations within SKSU campus boundaries.",
        "To provide staff with comprehensive privacy controls including a 'Visibility Toggle' to control when they are trackable.",
        "To reduce the time students spend searching for staff members by at least 60%.",
        "To implement geofencing technology that automatically disables tracking once a staff member leaves the university perimeter.",
        "To create a searchable directory system that allows students to filter staff by Department, Name, or Current Availability.",
        "To develop an administrative module for department management and anonymized analytics.",
        "To ensure full compliance with data privacy regulations including the Data Privacy Act of 2012 (RA 10173).",
        "To achieve zero-cost implementation using only free and open-source technologies."
    ]
    
    for i, obj in enumerate(objectives, 1):
        add_numbered_item(doc, i, obj)
    
    doc.add_page_break()
    
    # ==================== 5. SCOPE AND LIMITATIONS ====================
    create_styled_heading(doc, "5. Scope and Limitations", level=1)
    
    create_styled_heading(doc, "5.1 Project Scope", level=2)
    
    scope_items = ["Development of cross-platform mobile applications for Android and iOS devices", "Integration with OpenStreetMap and Flutter Map for free, open-source location visualization", "Implementation of Firebase free tier backend for real-time data synchronization", "Development of three user modules: Staff, Student, and Administrative", "Implementation of geofencing for SKSU campus boundaries", "Integration with university email system for authentication", "Development of privacy controls and consent mechanisms", "Zero-cost deployment using free and open-source technologies only"]
    for item in scope_items:
        add_bullet_point(doc, item)
    
    create_styled_heading(doc, "5.2 Limitations", level=2)
    
    limitations = ["The system requires active internet connectivity for real-time tracking", "GPS accuracy may vary depending on device hardware and environmental factors", "Initial deployment will be limited to the SKSU Main Campus", "The system does not track indoor floor levels (single-plane tracking only)", "Battery consumption may increase on staff devices when tracking is enabled", "The system requires voluntary participation from faculty and staff"]
    for item in limitations:
        add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ==================== 6. PROPOSED FEATURES ====================
    create_styled_heading(doc, "6. Proposed Features", level=1)
    
    create_styled_heading(doc, "6.1 Staff Module (The 'Beacon')", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("The Staff Module serves as the location broadcasting component of the system, giving faculty and staff complete control over their visibility.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    staff_features = [
        ("Privacy Toggle: ", "A master ON/OFF switch for location broadcasting. Staff can instantly enable or disable their visibility with a single tap."),
        ("Status Presets: ", "Manual status updates including: 'Available for Consultation,' 'In a Class,' 'In a Meeting,' 'Break Time,' 'Office Hours,' and 'Do Not Disturb.'"),
        ("Availability Status: ", "Rich availability indicators (Available, Busy, In Meeting, Teaching, On Break, Out of Office, Do Not Disturb) with color-coded display visible to students in real-time."),
        ("Auto-Kill Timer: ", "Automatically turns off tracking after office hours based on customizable schedule settings."),
        ("Quick Messages: ", "Pre-set messages that can be broadcast to searching students (e.g., 'Back in 10 minutes,' 'See me tomorrow')."),
        ("Location Override: ", "Ability to manually set a static location when GPS signal is weak (e.g., 'Currently at Admin Building')."),
        ("Schedule Integration: ", "Optional sync with class schedules to automatically update status during teaching hours."),
        ("Custom Status Message: ", "Add personalized messages to your status (e.g., 'Available until 3PM', 'In room 201').")
    ]
    
    for title, desc in staff_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    create_styled_heading(doc, "6.2 Student Module (The 'Seeker')", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("The Student Module provides the search and visualization interface for locating faculty and staff.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    student_features = [
        ("Searchable Directory: ", "Filter staff by Department, Name, Subject Taught, or Current Availability status."),
        ("Live Map View: ", "Real-time map showing moving markers (avatars) of online staff members with smooth animations using OpenStreetMap."),
        ("One-Tap Navigation: ", "Get walking directions from the student's current location to the staff member using open-source routing."),
        ("Favorites List: ", "Save frequently consulted faculty members for quick access."),
        ("Notification Alerts: ", "Opt-in notifications when a specific faculty member comes online or becomes available."),
        ("Estimated Walking Time: ", "Display estimated time to reach the faculty member based on current distance."),
        ("Office Hours Display: ", "View faculty consultation hours and schedule information."),
        ("Offline Mode: ", "Access cached faculty data and campus maps even without internet connection using SQLite local storage."),
        ("Push Notifications: ", "Receive important announcements and updates directly on your device."),
        ("Onboarding Tutorial: ", "Interactive guided tour for new users to learn all app features quickly."),
        ("In-App Updates: ", "Automatic notification and download of new app versions without visiting external stores.")
    ]
    
    for title, desc in student_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    create_styled_heading(doc, "6.3 Administrative Module", level=2)
    
    admin_features = [
        ("Department Management: ", "Manage the list of authorized faculty and staff accounts, assign departments, and control access levels."),
        ("User Verification: ", "Approve or reject registration requests from faculty and staff."),
        ("Analytics Dashboard: ", "View anonymized heatmaps of student-staff interaction points and peak consultation times."),
        ("Campus Boundary Management: ", "Define and update the geofencing polygon coordinates for campus boundaries."),
        ("System Configuration: ", "Configure default privacy settings, tracking intervals, and notification preferences."),
        ("Report Generation: ", "Generate usage reports for administrative review and system improvement.")
    ]
    
    for title, desc in admin_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    doc.add_page_break()
    
    # ==================== 7. APPLICATION UI DESIGN ====================
    create_styled_heading(doc, "7. Application User Interface Design", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("The following wireframes and mockups illustrate the proposed user interface design for UniTrack. The design follows modern mobile UI/UX principles with a focus on simplicity, accessibility, and intuitive navigation.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    # Login Screen
    create_styled_heading(doc, "7.1 Login Screen", level=2)
    para = doc.add_paragraph()
    run = para.add_run("The login screen provides secure authentication using SKSU email credentials. Users can choose their role (Student or Faculty) for appropriate access.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/01_login_screen.png", "Figure 7.1: UniTrack Login Screen")
    
    # Staff Dashboard
    create_styled_heading(doc, "7.2 Staff Dashboard (Beacon Module)", level=2)
    para = doc.add_paragraph()
    run = para.add_run("The Staff Dashboard features a prominent privacy toggle, status selection, and quick message options. Faculty members have full control over their visibility and can set their availability status with one tap.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/02_staff_dashboard.png", "Figure 7.2: Staff Dashboard with Privacy Controls")
    
    doc.add_page_break()
    
    # Student Directory
    create_styled_heading(doc, "7.3 Faculty Directory (Seeker Module)", level=2)
    para = doc.add_paragraph()
    run = para.add_run("Students can search and filter faculty members by name, department, or availability status. Each listing shows the faculty member's current status with color-coded indicators for quick identification.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/03_student_directory.png", "Figure 7.3: Faculty Directory with Search and Filters")
    
    # Live Map
    create_styled_heading(doc, "7.4 Live Campus Map", level=2)
    para = doc.add_paragraph()
    run = para.add_run("The live map displays the SKSU campus with building overlays and real-time faculty location markers. Markers are color-coded by availability status, and students can tap on any marker to view details and get directions.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/04_live_map.png", "Figure 7.4: Live Campus Map with Faculty Locations")
    
    doc.add_page_break()
    
    # Navigation
    create_styled_heading(doc, "7.5 Navigation & Directions", level=2)
    para = doc.add_paragraph()
    run = para.add_run("When a student selects a faculty member, the app provides turn-by-turn walking directions with estimated arrival time. The route is displayed on the map with clear visual guidance.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/05_navigation.png", "Figure 7.5: Navigation Screen with Walking Directions")
    
    # Privacy Settings
    create_styled_heading(doc, "7.6 Privacy Settings", level=2)
    para = doc.add_paragraph()
    run = para.add_run("Faculty members have granular control over their privacy settings, including auto-off timers, campus-only restrictions, and message preferences. All settings are clearly labeled with toggle switches.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/06_privacy_settings.png", "Figure 7.6: Privacy Settings Panel")
    
    doc.add_page_break()
    
    # Admin Dashboard
    create_styled_heading(doc, "7.7 Admin Dashboard", level=2)
    para = doc.add_paragraph()
    run = para.add_run("The administrative dashboard provides an overview of system usage, active faculty, student activity statistics, and quick access to management functions.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=8, line_spacing=1.5)
    
    add_mockup_image(doc, "mockups/07_admin_dashboard.png", "Figure 7.7: Administrative Dashboard")
    
    # Design Principles
    create_styled_heading(doc, "7.8 Design Principles", level=2)
    
    principles = [
        ("Color Scheme: ", "Green (#009933) for online/available status, Orange for busy, Red for unavailable, and Dark Blue (#003366) for headers and navigation."),
        ("Typography: ", "Clean, readable fonts with clear hierarchy. Large touch targets for mobile usability."),
        ("Privacy Indicators: ", "Prominent visual feedback when location sharing is active, with clear ON/OFF states."),
        ("Accessibility: ", "High contrast colors, readable font sizes, and support for screen readers."),
        ("Responsive Design: ", "Adapts to different screen sizes across Android and iOS devices.")
    ]
    
    for title, desc in principles:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    doc.add_page_break()
    
    # ==================== 8. TECHNICAL STACK ====================
    create_styled_heading(doc, "8. Technical Stack (Free & Open-Source)", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack is built entirely on free and open-source technologies, ensuring ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = para.add_run("zero licensing costs")
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 128, 0)
    run3 = para.add_run(" while maintaining professional-grade quality and scalability:")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    tech_headers = ['Component', 'Technology', 'Justification']
    tech_data = [
        ('Frontend', 'Flutter (Dart)', 'Free, open-source cross-platform development'),
        ('Map Engine', 'OpenStreetMap + Flutter Map', 'Free mapping with no API costs'),
        ('Backend/Database', 'Firebase Spark (Free Tier)', 'Free real-time NoSQL database'),
        ('Authentication', 'Firebase Auth (Free)', 'Free unlimited user authentication'),
        ('Location Services', 'Geolocator Package', 'Free Flutter GPS tracking package'),
        ('Hosting', 'Firebase Hosting (Free)', 'Free hosting with SSL included')
    ]
    
    create_styled_table(doc, tech_headers, tech_data)
    
    doc.add_paragraph()
    
    create_styled_heading(doc, "8.1 Development Tools (All Free/Open-Source)", level=2)
    
    tools = ["IDE: Visual Studio Code (Free, Open-Source) with Flutter extensions", "Version Control: Git with GitHub Free repositories", "UI/UX Design: Figma Free Tier (3 projects)", "Project Management: Trello Free Tier / GitHub Projects", "Testing: Flutter Test Framework (built-in), Android Emulator (free)", "CI/CD: GitHub Actions Free Tier (2,000 minutes/month)"]
    for tool in tools:
        add_bullet_point(doc, tool)
    
    doc.add_page_break()
    
    # ==================== 9. SYSTEM ARCHITECTURE ====================
    create_styled_heading(doc, "9. System Architecture", level=1)
    
    create_styled_heading(doc, "9.1 High-Level Architecture", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack follows a client-server architecture with real-time synchronization capabilities:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    arch_steps = [
        ("Data Acquisition: ", "The Staff mobile app retrieves GPS coordinates via the device's location services."),
        ("Privacy Check: ", "Before transmission, the app verifies that the Privacy Toggle is ON and the device is within the geofenced campus boundary."),
        ("Transmission: ", "If conditions are met, coordinates are pushed to Firebase Firestore via secure HTTPS connection."),
        ("Cloud Processing: ", "Firebase processes incoming data, validates geofence boundaries, and triggers notifications."),
        ("Synchronization: ", "Student apps 'listen' to Firebase Firestore streams using real-time listeners."),
        ("Rendering: ", "The Flutter Map package with OpenStreetMap renders the movement of markers on the student's screen.")
    ]
    
    for i, (title, desc) in enumerate(arch_steps, 1):
        p = doc.add_paragraph()
        run1 = p.add_run(f"{i}. ")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        run1.font.color.rgb = RGBColor(0, 51, 102)
        run2 = p.add_run(title)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.bold = True
        run3 = p.add_run(desc)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(12)
        set_paragraph_spacing(p, before=4, after=8, line_spacing=1.5)
        p.paragraph_format.left_indent = Inches(0.25)
    
    create_styled_heading(doc, "9.2 Database Schema", level=2)
    
    collections = [
        ("users: ", "Stores user profiles, roles, department affiliations, and preferences"),
        ("locations: ", "Real-time location data with coordinates and timestamp (overwritten, not logged)"),
        ("departments: ", "Department information and authorized personnel"),
        ("geofences: ", "Campus boundary polygon coordinates"),
        ("analytics: ", "Anonymized interaction data for reporting")
    ]
    
    for name, desc in collections:
        add_bullet_point(doc, desc, bold_prefix=name)
    
    doc.add_page_break()
    
    # ==================== 10. PRIVACY & SECURITY ====================
    create_styled_heading(doc, "10. Privacy & Security Measures", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack is designed with privacy as a foundational principle, ensuring compliance with the Data Privacy Act of 2012 (Republic Act No. 10173).")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    create_styled_heading(doc, "10.1 Core Privacy Principles", level=2)
    
    privacy_items = [
        ("Opt-in Only: ", "Location tracking is disabled by default. Staff must explicitly enable tracking."),
        ("No History Logging: ", "The system only stores the current coordinate. Past locations are immediately overwritten."),
        ("Campus Boundary Enforcement: ", "Geofencing automatically stops tracking outside the SKSU campus."),
        ("Minimal Data Collection: ", "Only essential data is collected. No behavioral tracking."),
        ("User Control: ", "Staff can disable tracking at any time with immediate effect.")
    ]
    
    for title, desc in privacy_items:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    create_styled_heading(doc, "10.2 Security Measures", level=2)
    
    security_items = [
        ("Encrypted Transmission: ", "All data uses TLS 1.3 encryption via HTTPS."),
        ("Authentication: ", "Firebase Authentication with university email verification."),
        ("Role-Based Access Control: ", "Different access levels for students, staff, and administrators."),
        ("Audit Logging: ", "Administrative actions are logged for accountability.")
    ]
    
    for title, desc in security_items:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    doc.add_page_break()
    
    # ==================== 11. DEVELOPMENT ROADMAP ====================
    create_styled_heading(doc, "11. Development Roadmap", level=1)
    
    phases = [
        ("Phase 1: Planning and Design (Weeks 1-4)", ["Requirements gathering and stakeholder interviews", "System architecture design", "UI/UX wireframe and prototype development", "Database schema design", "Privacy impact assessment"]),
        ("Phase 2: Core Development (Weeks 5-12)", ["Firebase project setup", "User authentication module", "Staff module with privacy controls", "Student module with map integration", "Real-time location synchronization", "Geofencing logic"]),
        ("Phase 3: Feature Enhancement (Weeks 13-16)", ["Administrative module", "Analytics dashboard", "Notification system", "Navigation feature", "Status presets and messaging"]),
        ("Phase 4: Testing and QA (Weeks 17-20)", ["Unit testing", "Integration testing", "User acceptance testing with SKSU users", "Performance and security testing", "Bug fixes and optimization"]),
        ("Phase 5: Deployment and Launch (Weeks 21-24)", ["APK distribution setup", "Production environment", "User training", "Pilot launch", "Full campus rollout"])
    ]
    
    for phase_title, items in phases:
        h = doc.add_paragraph()
        run = h.add_run(phase_title)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 76, 153)
        set_paragraph_spacing(h, before=12, after=6)
        
        for item in items:
            add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ==================== 12. PROJECT TIMELINE ====================
    create_styled_heading(doc, "12. Project Timeline", level=1)
    
    doc.add_paragraph()
    
    timeline_headers = ['Phase', 'Duration', 'Start', 'End', 'Deliverables']
    timeline_data = [
        ('Planning & Design', '4 weeks', 'Week 1', 'Week 4', 'SRS, Wireframes'),
        ('Core Development', '8 weeks', 'Week 5', 'Week 12', 'Working Modules'),
        ('Feature Enhancement', '4 weeks', 'Week 13', 'Week 16', 'Complete Features'),
        ('Testing & QA', '4 weeks', 'Week 17', 'Week 20', 'Test Reports'),
        ('Deployment', '4 weeks', 'Week 21', 'Week 24', 'Live Application')
    ]
    
    create_styled_table(doc, timeline_headers, timeline_data)
    
    doc.add_paragraph()
    
    total = doc.add_paragraph()
    total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = total.add_run("Total Project Duration: 24 weeks (approximately 6 months)")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_page_break()
    
    # ==================== 13. COST ANALYSIS ====================
    create_styled_heading(doc, "13. Cost Analysis: Zero-Budget Implementation", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("One of the key advantages of UniTrack is that it can be developed and deployed entirely using free and open-source technologies, requiring no financial investment.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    budget_headers = ['Component', 'Free Solution', 'Cost']
    budget_data = [
        ('Development IDE', 'Visual Studio Code (Open-Source)', 'FREE'),
        ('Map Engine', 'OpenStreetMap + Flutter Map', 'FREE'),
        ('Backend Database', 'Firebase Firestore (Spark Free Plan)', 'FREE'),
        ('Authentication', 'Firebase Auth (Unlimited users)', 'FREE'),
        ('Cloud Hosting', 'Firebase Hosting (10GB storage)', 'FREE'),
        ('UI/UX Design', 'Figma Free Tier', 'FREE'),
        ('Version Control', 'GitHub Free', 'FREE'),
        ('CI/CD Pipeline', 'GitHub Actions (2,000 min/month)', 'FREE'),
        ('App Distribution', 'Direct APK / GitHub Releases', 'FREE')
    ]
    
    create_styled_table(doc, budget_headers, budget_data, header_color='006400')
    
    doc.add_paragraph()
    
    total_box = doc.add_paragraph()
    total_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = total_box.add_run("TOTAL PROJECT COST: PHP 0.00")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # ==================== 14. RISK ASSESSMENT ====================
    create_styled_heading(doc, "14. Risk Assessment and Mitigation", level=1)
    
    doc.add_paragraph()
    
    risk_headers = ['Risk', 'Probability', 'Impact', 'Mitigation Strategy']
    risk_data = [
        ('Low faculty adoption', 'Medium', 'High', 'Awareness campaigns; privacy emphasis'),
        ('GPS accuracy issues', 'Medium', 'Medium', 'Location smoothing; manual override'),
        ('Battery drain', 'High', 'Medium', 'Optimized polling frequency'),
        ('Privacy concerns', 'Medium', 'High', 'Transparent policy; user control'),
        ('Technical failures', 'Low', 'High', 'Redundant systems; backups'),
        ('Firebase limits exceeded', 'Low', 'Medium', 'Monitor usage; optimize queries')
    ]
    
    create_styled_table(doc, risk_headers, risk_data, header_color='8B0000')
    
    doc.add_page_break()
    
    # ==================== 15. EXPECTED OUTCOMES ====================
    create_styled_heading(doc, "15. Expected Outcomes", level=1)
    
    outcomes = [
        ("Improved Time Management: ", "Students will save an estimated 60% of time previously spent searching for faculty."),
        ("Reduced Physical Traffic: ", "Faculty rooms will experience reduced congestion."),
        ("Enhanced Communication: ", "Real-time status updates improve student-faculty communication."),
        ("Modernized Campus: ", "SKSU positioned as a tech-forward institution."),
        ("Data-Driven Insights: ", "Anonymized analytics help optimize consultation hours."),
        ("Zero-Cost Implementation: ", "Complete project delivery without financial investment.")
    ]
    
    for title, desc in outcomes:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    # ==================== 16. EVALUATION CRITERIA ====================
    create_styled_heading(doc, "16. Evaluation Criteria", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("Success will be measured against the following KPIs:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    
    kpis = [
        ("User Adoption Rate: ", "70% faculty and 80% students within first semester"),
        ("Time Savings: ", "60% reduction in time spent locating faculty"),
        ("System Uptime: ", "99.5% availability during operational hours"),
        ("User Satisfaction: ", "4.0/5.0 average rating"),
        ("Privacy Compliance: ", "Zero privacy incidents or data breaches"),
        ("Response Time: ", "Location updates visible within 3 seconds")
    ]
    
    for title, desc in kpis:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    doc.add_page_break()
    
    # ==================== 17. FUTURE EXPANSION: CAMPUS KIOSK INTEGRATION ====================
    create_styled_heading(doc, "17. Future Expansion: Campus Kiosk Integration", level=1)
    
    para = doc.add_paragraph()
    run = para.add_run("UniTrack is designed with scalability in mind, allowing for future deployment on campus kiosk displays. This expansion would bring the existing mobile application features to large-screen installations, providing the same functionality accessible to students who may not have their phones readily available.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=12, line_spacing=1.5)
    para.paragraph_format.first_line_indent = Inches(0.5)
    
    create_styled_heading(doc, "17.1 Kiosk Display Features (Mirroring Mobile App)", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("Campus kiosks would display the same real-time data already available in the UniTrack mobile application:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=8, line_spacing=1.5)
    
    kiosk_features = [
        ("Faculty Directory: ", "Searchable list of all faculty members with their department, profile photo, and contact information—same as the mobile app's directory screen."),
        ("Real-Time Availability Status: ", "Display of each faculty member's current status (Available, Busy, In Meeting, Teaching, On Break, Out of Office, Do Not Disturb) as set in the mobile app."),
        ("Custom Status Messages: ", "Faculty-written status messages explaining their availability, such as 'In consultation until 3PM' or 'Available for walk-ins.'"),
        ("Live Campus Map: ", "Interactive map showing faculty locations on campus, identical to the mobile app's MapLibre-powered map view."),
        ("Walking Directions: ", "Step-by-step navigation from the kiosk location to any faculty member's office, using the same routing as the mobile app."),
        ("Department Filtering: ", "Filter faculty by department to quickly find relevant staff, same functionality as the mobile app."),
        ("Estimated Walking Time: ", "Display distance and estimated walking time to each faculty member's location.")
    ]
    
    for title, desc in kiosk_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    create_styled_heading(doc, "17.2 Kiosk Placement Recommendations", level=2)
    
    placement_features = [
        ("Campus Entrances: ", "Main gates and building lobbies where students and visitors first arrive on campus."),
        ("Department Offices: ", "Common areas outside department clusters for students looking for specific faculty."),
        ("Student Centers: ", "High-traffic areas like canteens, libraries, and study halls."),
        ("Administration Building: ", "Near registrar, guidance, and other administrative offices.")
    ]
    
    for title, desc in placement_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    create_styled_heading(doc, "17.3 Technical Implementation", level=2)
    
    para = doc.add_paragraph()
    run = para.add_run("The kiosk system would leverage the existing UniTrack infrastructure:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=6, after=8, line_spacing=1.5)
    
    tech_features = [
        ("Same Firebase Backend: ", "Kiosks connect to the same Firebase Firestore database used by the mobile app, ensuring real-time data synchronization."),
        ("Web-Based Interface: ", "Kiosks run the Flutter web version of UniTrack, requiring only a browser and internet connection."),
        ("No Additional Server Costs: ", "Firebase's existing infrastructure handles kiosk requests within the current free tier or minimal usage fees."),
        ("Offline Mode Support: ", "Kiosks can cache faculty data locally, same as the mobile app's offline functionality.")
    ]
    
    for title, desc in tech_features:
        add_bullet_point(doc, desc, bold_prefix=title)
    
    # Cost estimate for kiosk expansion
    doc.add_paragraph()
    
    kiosk_headers = ['Component', 'Estimated Units', 'Estimated Cost']
    kiosk_data = [
        ('Touchscreen Kiosk (32")', '3-5 units', 'PHP 25,000-50,000 each'),
        ('Kiosk Stand/Enclosure', '3-5 units', 'PHP 5,000-10,000 each'),
        ('Internet Connection', 'Existing campus WiFi', 'PHP 0 (uses campus network)'),
        ('Software Development', 'Web deployment', 'PHP 0 (Flutter web already supported)'),
    ]
    
    create_styled_table(doc, kiosk_headers, kiosk_data, header_color='006400')
    
    note = doc.add_paragraph()
    run = note.add_run("Note: Kiosk expansion is optional and can be implemented based on budget availability. The kiosks display the same data as the mobile app—no additional features or hardware sensors are required. The core mobile application remains completely free.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    set_paragraph_spacing(note, before=8, after=12)
    
    doc.add_page_break()
    
    # ==================== 18. CONCLUSION ====================
    create_styled_heading(doc, "18. Conclusion", level=1)
    
    para1 = doc.add_paragraph()
    run = para1.add_run("UniTrack represents a significant step forward in modernizing campus operations at Sultan Kudarat State University. By leveraging free and open-source technologies, the system addresses student-faculty connectivity challenges while maintaining privacy—all at zero cost.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para1, before=6, after=12, line_spacing=1.5)
    para1.paragraph_format.first_line_indent = Inches(0.5)
    
    para2 = doc.add_paragraph()
    run = para2.add_run("The privacy-first design ensures faculty maintain complete control over their visibility with robust safeguards including opt-in consent, geofencing, and zero historical tracking.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para2, before=0, after=12, line_spacing=1.5)
    para2.paragraph_format.first_line_indent = Inches(0.5)
    
    para3 = doc.add_paragraph()
    run = para3.add_run("With a realistic 24-week timeline and completely free technology stack, UniTrack is positioned for successful implementation. The use of OpenStreetMap, Firebase free tier, and Flutter ensures professional-grade quality without licensing fees.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para3, before=0, after=12, line_spacing=1.5)
    para3.paragraph_format.first_line_indent = Inches(0.5)
    
    para4 = doc.add_paragraph()
    run = para4.add_run("UniTrack will not only solve the immediate problem of locating faculty but will establish SKSU as a leader in educational technology innovation in the SOCCSKSARGEN region and beyond.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para4, before=0, after=12, line_spacing=1.5)
    para4.paragraph_format.first_line_indent = Inches(0.5)
    
    # ==================== 19. REFERENCES ====================
    create_styled_heading(doc, "19. References", level=1)
    
    references = [
        "OpenStreetMap Foundation. (2026). OpenStreetMap Documentation. https://wiki.openstreetmap.org",
        "Flutter Map Package. (2026). https://pub.dev/packages/flutter_map",
        "Firebase Documentation. (2026). https://firebase.google.com/docs",
        "Flutter Documentation. (2026). https://docs.flutter.dev",
        "Republic Act No. 10173 - Data Privacy Act of 2012. Official Gazette of the Philippines.",
        "National Privacy Commission. (2024). Guidelines on Privacy Impact Assessment.",
        "Sultan Kudarat State University. (2026). SKSU Strategic Development Plan.",
        "Geolocator Flutter Package. (2026). https://pub.dev/packages/geolocator",
        "Internet of Things in Education: A Review. IEEE Access. (2024).",
        "Smart Campus Solutions: Best Practices for Digital Transformation. (2025)."
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run1 = p.add_run(f"[{i}] ")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run1.bold = True
        run1.font.color.rgb = RGBColor(0, 51, 102)
        run2 = p.add_run(ref)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        set_paragraph_spacing(p, before=2, after=4, line_spacing=1.15)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    # Save document
    doc.save('UniTrack_Project_Proposal_SKSU_Final.docx')
    print("Document created successfully: UniTrack_Project_Proposal_SKSU_Final.docx")
    return 'UniTrack_Project_Proposal_SKSU_Final.docx'

if __name__ == "__main__":
    create_proposal()
